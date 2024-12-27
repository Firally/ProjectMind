# RAG_service.py

import os
import sqlite3
import torch
from fastapi import FastAPI, HTTPException, File, UploadFile, BackgroundTasks, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from langchain.schema import Document
from langchain_community.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv
from pydantic import BaseModel
from openai import OpenAI
from typing import Optional
import tiktoken

import asyncio  # Для асинхронной работы
import PyPDF2  # Для обработки PDF файлов
from docx import Document as DocxDocument  # Для обработки DOCX файлов

load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI(api_key=openai_api_key)

def ask_openai(context, query, model: str = "gpt-3.5-turbo", max_tokens: int = 16000):
    messages = [
        {"role": "system", "content": "Ты помощник, который отвечает на вопросы максимально точно и кратко, используя предоставленную информацию."},
        {"role": "user", "content": f"У тебя есть следующая информация:\n{context}\n\nВопрос: {query}"}
    ]
    
    total_tokens = sum(count_tokens(message["content"], model) for message in messages)
    
    if total_tokens > max_tokens:
        allowed_tokens = max_tokens - count_tokens(messages[0]["content"], model) - count_tokens(messages[1]["content"], model)
        encoding = tiktoken.encoding_for_model(model)
        encoded_context = encoding.encode(context)
        truncated_context = encoding.decode(encoded_context[:allowed_tokens])
        messages[1]["content"] = f"У тебя есть следующая информация:\n{truncated_context}\n\nВопрос: {query}"
    
    response = client.chat.completions.create(
        model=model,
        messages=messages,
    )
    return response

conn = sqlite3.connect("projects_with_files.db")
cursor = conn.cursor()
embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

app = FastAPI()

origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "https://a11f-57-129-59-152.ngrok-free.app",
    "http://192.168.31.154:3000"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

class QueryRequest(BaseModel):
    user_id: str
    query: str
    top_k: int

create_table_query = """
CREATE TABLE IF NOT EXISTS project_tasks (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_id TEXT NOT NULL,
    project_name TEXT NOT NULL,
    task_name TEXT NOT NULL,
    note_text TEXT,
    file_path TEXT,
    file_content BLOB
);
"""
cursor.execute(create_table_query)
conn.commit()

class ProjectTask(BaseModel):
    user_id: str
    project_name: str
    task_name: str
    note_text: str

@app.post("/create_task")
async def create_task(
    background_tasks: BackgroundTasks,
    user_id: str = Form(...),
    project_name: str = Form(...),
    task_name: str = Form(...),
    note_text: str = Form(...),
    file_path: str = Form(None),
    file: Optional[UploadFile] = File(None)
):
    try:
        print(f"Полученные данные: user_id={user_id}, project_name={project_name}, task_name={task_name}, note_text={note_text}")

        file_content = None

        if file:
            upload_dir = "/home/baton/Documents/frontend_project/uploads"
            os.makedirs(upload_dir, exist_ok=True)
            file_location = os.path.join(upload_dir, file.filename)

            with open(file_location, "wb") as f:
                f.write(await file.read())

            file_path = file_location

            with open(file_location, "rb") as f:
                file_content = f.read()
        elif file_path:
            with open(file_path, "rb") as f:
                file_content = f.read()

        insert_query = """
        INSERT INTO project_tasks (user_id, project_name, task_name, note_text, file_path, file_content)
        VALUES (?, ?, ?, ?, ?, ?);
        """
        cursor.execute(insert_query, (
            user_id,
            project_name,
            task_name,
            note_text,
            file_path,
            file_content
        ))
        conn.commit()

        background_tasks.add_task(update_faiss_index_func, user_id)

        return {"message": "Задача успешно создана."}
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/get_relevant")
async def get_relevant(request: QueryRequest):
    try:
        vectorstore = FAISS.load_local(f"faiss_index_{request.user_id}", embeddings, allow_dangerous_deserialization=True)
        results = vectorstore.similarity_search(request.query, k=request.top_k)
        response = "\n\n".join([res.page_content for res in results])
        return {"text": response}
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/get_LLM_answer")
async def get_LLM_answer(request: QueryRequest):
    try:
        relevant_response = await get_relevant(request)
        context = relevant_response['text']
        response = ask_openai(context, request.query)
        return {"model_answer": response.choices[0].message.content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class UpdateFAISSRequest(BaseModel):
    user_id: str

@app.post("/update_faiss_index")
async def update_faiss_index(request: UpdateFAISSRequest, background_tasks: BackgroundTasks):
    user_id = request.user_id
    background_tasks.add_task(update_faiss_index_func, user_id)
    return {"message": f"FAISS индекс для пользователя {user_id} обновляется в фоновом режиме."}

@app.delete("/delete_project")
async def delete_project(background_tasks: BackgroundTasks, user_id: str = Query(...), project_name: str = Query(...)):
    try:
        delete_query = """
        DELETE FROM project_tasks WHERE user_id = ? AND project_name = ?;
        """
        cursor.execute(delete_query, (user_id, project_name))
        conn.commit()
        background_tasks.add_task(update_faiss_index_func, user_id)
        return {"message": f"Проект '{project_name}' для пользователя {user_id} успешно удален."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

async def update_faiss_index_func(user_id: str):
    try:
        cursor.execute("SELECT id, project_name, task_name, note_text, file_path, file_content FROM project_tasks WHERE user_id = ?", (user_id,))
        rows = cursor.fetchall()

        documents = []
        for row in rows:
            record_id, project_name, task_name, note_text, file_path, file_content = row
            file_text = ""
            if file_content and file_path:
                temp_path = save_blob_to_temp(file_content, file_path)
                if file_path.endswith(".docx"):
                    doc = DocxDocument(temp_path)
                    file_text = "\n".join([p.text for p in doc.paragraphs])
                elif file_path.endswith(".pdf"):
                    with open(temp_path, "rb") as pdf_file:
                        reader = PyPDF2.PdfReader(pdf_file)
                        file_text = "\n".join([page.extract_text() if page.extract_text() else "" for page in reader.pages])
                else:
                    with open(temp_path, "r", encoding="utf-8") as file:
                        file_text = file.read()
                os.remove(temp_path)

            combined_text = f"Проект: {project_name}\nЗадача: {task_name}\nЗаметка: {note_text}\nСодержимое файла:\n{file_text}"
            documents.append(Document(page_content=combined_text, metadata={"id": record_id}))

        if not documents:
            print("Нет задач для обновления FAISS индекса.")
            return

        vectorstore = FAISS.from_documents(documents, embeddings)
        vectorstore.save_local(f"faiss_index_{user_id}")

        print(f"FAISS индекс для пользователя {user_id} успешно обновлен.")
    except Exception as e:
        print(f"Ошибка при обновлении FAISS индекса для пользователя {user_id}: {e}")

def save_blob_to_temp(blob_content, file_path):
    temp_dir = "/tmp"
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, os.path.basename(file_path))
    with open(temp_path, "wb") as file:
        file.write(blob_content)
    return temp_path

def count_tokens(text: str, model: str = "gpt-3.5-turbo") -> int:
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))